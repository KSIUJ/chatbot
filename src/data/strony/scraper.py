import time
import io
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse, unquote
from collections import deque, Counter
from pathlib import Path
import re

# wymagane dodatkowe biblioteki do obslugi plikow: pip install pypdf python-docx
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None


START_URLS = [
    "https://matinf.uj.edu.pl/",
    "https://kmsuj.matinf.uj.edu.pl/",
    "https://nkr.si/",
    "https://knmf.im.uj.edu.pl/",
]

ALLOWED_DOMAINS = {urlparse(u).netloc for u in START_URLS}
REPO_ROOT = Path(__file__).resolve().parents[3]
OUTPUT_FILE = REPO_ROOT / "data" / "strony" / "webiste_data.txt"
DELAY = 0.2
MAX_PAGES = 1000

WIKIPEDIA_URLS = [
    "https://pl.wikipedia.org/wiki/Wydzia%C5%82_Matematyki_i_Informatyki_Uniwersytetu_Jagiello%C5%84skiego",
]

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/150.0.0.0 Safari/537.36"
}

FALLBACK_USED = []
FALLBACK_FILES_USED = []  # pliki (pdf/docx/txt) ktorych nie dalo sie odczytac
SCRAPED_FILES = []  # pliki (pdf/docx/txt) ktore udalo sie zescrapowac

# klasy do pominiecia
NOISE_SELECTORS = [
    "[aria-hidden='true']",
    ".sr-only",
    ".hide-accessible",
    ".invisible-element",
    ".visually-hidden",
    ".visible-interaction",
    ".descArticles",
    ".skip-link",
    "#wmfc_quickAccessNav",
    "[id$='quickAccessNav']"
]

# frazy do pominiecia
LINE_STOPLIST = {
    "pomiń baner",
    "pomiń do treści",
    "wszystkie",
    "popularne",
    "artykuły",
    "ukryty",
    "menu",
    "zaloguj",
    "wersja kontrastowa",
    "a",
    "en",
    "hidden",
    "skip to content",
}

# zmienne do liczenia ilosci wstapienia frazy (jesli wystepuje na duzej ilosci stron to znaczy ze fragment boilerplatea)
BOILERPLATE_RATIO = 0.3
BOILERPLATE_MAX_WORDS = 6

# zmienne do usuwania menu ktore sie powialy na min ilosci stron
MENU_ITEM_MAX_WORDS = 4
MENU_ITEM_MIN_PAGES = 3

# od ilu slow linia liczy sie jako tresc
DEDUPE_MIN_WORDS = 8

MIN_CONTENT_WORDS = 25
MIN_CONTENT_WORDS_FILES = 40

# rozszerzenia plikow ktore probujemy scrapowac jako dokumenty (nie html)
FILE_EXTENSIONS = {".pdf", ".docx", ".txt"}

# mapowanie Content-Type -> rozszerzenie, przydatne gdy url nie ma rozszerzenia
# (systemy typu Liferay czesto serwuja pliki pod adresami bez .pdf/.docx w sciezce)
CONTENT_TYPE_EXTENSIONS = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "text/plain": ".txt",
}

# pliki dopuszczone do scrapowania - tylko te ktore pasuja do ponizszych slow kluczowych
WHITELIST_FILE_KEYWORDS = [
    "poradnik_pierwszaka", "poradnik pierwszaka",
    "reg-dypl",
    "regulamin+dypl", "regulamin_dypl",
    "minima-smp",
    "zagadnien",
    "egz-lic",
    "egz-mgr", "egz_mgr",
    "egzamin-licencjacki",
    "egzamin-magisterski",
    "programy-studiow",
    "harmonogram_czynnosci",
    "harmonogram-czynnosci",
    "wniosek",
    "formularz",
    "podanie",
    "urlop",
    "wpis+warunkowy", "wpis_warunkowy",
    "wpis+po+urlopie", "wpis_po_urlopie",
    "powtarzanie",
    "rezygnacja",
    "przepisanie",
    "uznanie_przedmiotow", "uznanie przedmiotow",
    "dopisanie",
    "duplikat",
    "dolaczenie", "usuniecie",
    "zaliczenie",
    "rozszerzenie+programu", "rozszerzenie_programu",
    "awans_wpis",
    "stypendium",
    "stypendia",
    "kalendarz",
]

WHITELIST_EXCLUDE_OVERRIDE = [
    "lista", "laureat", "rankingow", "wynik", "stypendysci", "stypendyści",
    "minigrant", "tutoring", "konferenc", "wyjazd", "wyjazdy",
    "dofinansowanie", "nagrody_za_grant", "nagroda_za_grant",
    "research support", "research_support", "doktorant",
    "bieg", "charytatyw", "archiwizacj", "kalendarz_2026_web", "kalendarz-2026-web",
]


def is_whitelisted_file(url: str) -> bool:
    decoded = unquote(url).lower()
    if any(bad in decoded for bad in WHITELIST_EXCLUDE_OVERRIDE):
        return False
    return any(good in decoded for good in WHITELIST_FILE_KEYWORDS)


def is_same_domain(url: str) -> bool:
    return urlparse(url).netloc in ALLOWED_DOMAINS


def clean_url(url: str) -> str:
    parsed = urlparse(url)
    return parsed._replace(
        scheme="https",
        netloc=parsed.netloc.lower(),
        fragment="",
        query="",
    ).geturl()

# usuwa niepotrzebne jezyki np /en/
LOCALE_PATTERN = re.compile(r"^/[a-z]{2}(_[a-z]{2})?/", re.IGNORECASE)
SKIP_PATTERNS = ("/c/portal/login", "/c/portal/logout")

def should_skip(url: str) -> bool:
    path = urlparse(url).path.lower()
    if any(pattern in path for pattern in SKIP_PATTERNS):
        return True
    if LOCALE_PATTERN.match(path):
        return True
    # powtarzajacy sie segment journal_content (linkuje do siebie w kolo)
    if path.count("/journal_content/") > 1:
        return True
    return False


def detect_file_ext(url: str, content_type: str):
    # rozpoznajemy typ pliku najpierw po Content-Type (bardziej niezawodne niz url),
    # bo systemy typu Liferay czesto serwuja pliki pod adresami bez rozszerzenia
    content_type = content_type.split(";")[0].strip().lower()
    if content_type in CONTENT_TYPE_EXTENSIONS:
        return CONTENT_TYPE_EXTENSIONS[content_type]
    # fallback - rozszerzenie na koncu sciezki url
    path = urlparse(url).path.lower()
    for ext in FILE_EXTENSIONS:
        if path.endswith(ext):
            return ext
    return None


def is_file_link(url: str) -> bool:
    # rozpoznaje link do pliku (pdf/docx/txt) po rozszerzeniu w ścieżce URL
    path = urlparse(url).path.lower()
    segments = path.split("/")
    return any(seg.endswith(ext) for seg in segments for ext in FILE_EXTENSIONS)


# elementy blokowe (do tworzenia spojnych akapitow)
BLOCK_TAGS = ["p", "li", "h1", "h2", "h3", "h4", "h5", "h6", "td", "th", "caption", "dd", "dt", "blockquote", "figcaption"]


DOT_LEADER_PATTERN = re.compile(r"[.…]{3,}")  # kropki-wypelniacze z formularzy (np. "..................")

def _clean_spacing(text: str) -> str:
    text = DOT_LEADER_PATTERN.sub(" ", text)
    text = " ".join(text.split())
    text = re.sub(r"\s+([.,;:!?])", r"\1", text)
    text = re.sub(r"([„«])\s+", r"\1", text)
    text = re.sub(r"\s+([”»])", r"\1", text)
    return text

def extract_tables_as_sentences(content):
    lines = []
    for table in content.find_all("table"):
        rows = table.find_all("tr")
        if not rows:
            continue

        # ustal naglowki: albo z <th> w pierwszym wierszu, albo z pierwszego wiersza <td>
        header_cells = rows[0].find_all(["th", "td"])
        headers = [_clean_spacing(c.get_text(separator=" ", strip=True)) for c in header_cells]
        has_header = bool(rows[0].find_all("th")) or all(h and not h.replace(",", "").replace(".", "").isdigit() for h in headers if h)

        data_rows = rows[1:] if has_header else rows

        for row in data_rows:
            cells = row.find_all(["td", "th"])
            values = [_clean_spacing(c.get_text(separator=" ", strip=True)) for c in cells]
            if not any(values):
                continue

            if has_header and len(headers) == len(values):
                parts = [f"{h}: {v}" for h, v in zip(headers, values) if v]
            else:
                parts = [v for v in values if v]

            if parts:
                lines.append(", ".join(parts))

        table.decompose()  # usuwamy z drzewa, zeby generic get_block_text nie zdublowal komorek

    return lines

def get_block_text(content):
    # wyciaga tekst blok po bloku (p/li/h*) zeby link w srodku zdania go nie rozrywal
    lines = []
    for el in content.find_all(BLOCK_TAGS):
        if el.find_parent(BLOCK_TAGS) is not None:
            continue
        text = el.get_text(separator=" ", strip=True)
        text = _clean_spacing(text)
        if text:
            lines.append(text)

    if not lines:
        text = content.get_text(separator="\n")
        lines = [_clean_spacing(l) for l in text.splitlines() if l.strip()]

    return lines


def dedupe_repeated_blocks(lines):
    # usuwa caly blok linii, jesli powtarza sie bezposrednio po sobie (mobile+desktop menu)
    n = len(lines)
    result = []
    i = 0
    while i < n:
        matched = False
        max_block = (n - i) // 2
        for block_len in range(max_block, 0, -1):
            if lines[i:i + block_len] == lines[i + block_len:i + 2 * block_len]:
                result.extend(lines[i:i + block_len])
                i += 2 * block_len
                matched = True
                break
        if not matched:
            result.append(lines[i])
            i += 1
    return result


def strip_noise(content):
    # usuwa niepotrzebne elementy strony
    for tag in content(["script", "style", "noscript", "nav", "header", "footer"]):
        tag.decompose()
    for selector in NOISE_SELECTORS:
        for el in content.select(selector):
            el.decompose()
    return content


def extract_text(soup: BeautifulSoup):
    content = (
        soup.find(id="main-content")
        or soup.find("main")
        or soup.find("article")
        or soup.find("body")
    )
    if content is None:
        return None

    content = strip_noise(content)

    table_lines = extract_tables_as_sentences(content)   # <-- nowe, PRZED get_block_text
    lines = get_block_text(content)
    lines = table_lines + lines                          # albo wstaw w odpowiednim miejscu, jesli zalezy Ci na kolejnosci

    lines = [line for line in lines if line.lower() not in LINE_STOPLIST]
    lines = dedupe_repeated_blocks(lines)

    deduped = []
    for line in lines:
        if not deduped or deduped[-1] != line:
            deduped.append(line)

    return "\n".join(deduped) if deduped else None


def extract_wikipedia_text(soup: BeautifulSoup):
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    content = soup.find("div", id="mw-content-text")
    if content is None:
        return None

    # usuwa elementy niebedace trescia
    for selector in [
        ".mw-editsection", "sup.reference", ".reflist", ".navbox", ".ambox",
        ".infobox", "table.metadata", ".mw-references-wrap", "style",
        "#toc", ".toc", ".hatnote", "sup", ".mw-authority-control",
        ".noprint", ".mw-empty-elt",
    ]:
        for el in content.select(selector):
            el.decompose()

    lines = get_block_text(content)
    return "\n".join(lines)


def extract_pdf_bytes(content: bytes):
    # wyciaga tekst z pliku pdf strona po stronie
    if PdfReader is None:
        print("[BŁĄD - brak biblioteki pypdf, pomijam pliki PDF]")
        return None
    try:
        reader = PdfReader(io.BytesIO(content))
        pages_text = []
        for page in reader.pages:
            text = page.extract_text() or ""
            text = _clean_spacing(text)
            if text:
                pages_text.append(text)
        return "\n".join(pages_text) if pages_text else None
    except Exception as e:
        print(f"[BŁĄD - odczyt PDF] {e}")
        return None


def extract_docx_bytes(content: bytes):
    # wyciaga tekst z akapitow oraz tabel pliku docx
    if Document is None:
        print("[BŁĄD - brak biblioteki python-docx, pomijam pliki DOCX]")
        return None
    try:
        doc = Document(io.BytesIO(content))
        lines = []
        for para in doc.paragraphs:
            text = _clean_spacing(para.text)
            if text:
                lines.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = _clean_spacing(cell.text)
                    if text:
                        lines.append(text)
        return "\n".join(lines) if lines else None
    except Exception as e:
        print(f"[BŁĄD - odczyt DOCX] {e}")
        return None


def extract_txt_bytes(content: bytes):
    # probuje kilku kodowan, bo stare pliki uczelniane bywaja w cp1250
    text = None
    for encoding in ("utf-8", "cp1250", "iso-8859-2"):
        try:
            text = content.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        text = content.decode("utf-8", errors="ignore")

    lines = [_clean_spacing(l) for l in text.splitlines() if l.strip()]
    return "\n".join(lines) if lines else None


def scrape_file(url: str):
    # pobiera i wyciaga tresc z pliku (pdf/docx/txt), tylko jesli jest na whiteliscie
    if not is_whitelisted_file(url):
        return None

    try:
        resp = requests.get(url, headers=HEADERS, timeout=15)
        resp.raise_for_status()
    except Exception as e:
        print(f"[BŁĄD - plik] {url} -> {e}")
        return None

    ext = detect_file_ext(url, resp.headers.get("Content-Type", ""))
    if ext == ".pdf":
        text = extract_pdf_bytes(resp.content)
    elif ext == ".docx":
        text = extract_docx_bytes(resp.content)
    elif ext == ".txt":
        text = extract_txt_bytes(resp.content)
    else:
        text = None

    if text is None:
        FALLBACK_FILES_USED.append(url)
        print(f"[POMINIĘTO - brak treści pliku] {url}")
        return None

    SCRAPED_FILES.append(url)
    print(f"[PLIK] pobrano: {url}")
    return text


def remove_boilerplate(pages):
    # usuwa krotkie linie powtarzajace sie na wielu stronach
    total_pages = len(pages)
    if total_pages <= 1:
        return pages

    line_doc_count = Counter()
    for _, text in pages:
        unique_lines = set(text.splitlines())
        for line in unique_lines:
            line_doc_count[line] += 1

    threshold = max(2, int(total_pages * BOILERPLATE_RATIO))

    def is_boilerplate(line: str) -> bool:
        word_count = len(line.split())
        count = line_doc_count[line]
        if word_count <= BOILERPLATE_MAX_WORDS and count >= threshold:
            return True
        if word_count <= MENU_ITEM_MAX_WORDS and count >= MENU_ITEM_MIN_PAGES:
            return True
        return False

    cleaned_pages = []
    for url, text in pages:
        kept = [line for line in text.splitlines() if not is_boilerplate(line)]
        cleaned_pages.append((url, "\n".join(kept)))
    return cleaned_pages


def _normalize(line: str) -> str:
    return " ".join(line.lower().split())


def dedupe_content_across_pages(pages):
    # pozbywa sie duplikatow na wielu stronach
    line_pages = {}
    for url, text in pages:
        for line in text.splitlines():
            if len(line.split()) < DEDUPE_MIN_WORDS:
                continue
            key = _normalize(line)
            line_pages.setdefault(key, set()).add(url)

    canonical_url = {}
    for key, urls in line_pages.items():
        if len(urls) > 1:
            canonical_url[key] = max(urls, key=lambda u: len(urlparse(u).path))

    cleaned_pages = []
    for url, text in pages:
        kept = []
        for line in text.splitlines():
            key = _normalize(line) if len(line.split()) >= DEDUPE_MIN_WORDS else None
            if key in canonical_url and canonical_url[key] != url:
                continue  # duplikat - trzymamy go tylko na stronie kanonicznej
            kept.append(line)
        cleaned_pages.append((url, "\n".join(kept)))
    return cleaned_pages


def drop_thin_hub_pages(pages):
    kept, dropped = [], []
    for url, text in pages:
        word_count = len(text.split())
        threshold = MIN_CONTENT_WORDS_FILES if is_file_link(url) else MIN_CONTENT_WORDS
        if word_count < threshold:
            dropped.append((url, word_count))
        else:
            kept.append((url, text))
    return kept, dropped


def scrape_wikipedia():
    results = []
    for url in WIKIPEDIA_URLS:
        try:
            resp = requests.get(url, headers=HEADERS, timeout=15)
            resp.raise_for_status()
        except Exception as e:
            print(f"[BŁĄD - Wikipedia] {url} -> {e}")
            continue

        soup = BeautifulSoup(resp.text, "html.parser")
        page_text = extract_wikipedia_text(soup)

        if page_text is None:
            print(f"[POMINIĘTO - brak treści Wikipedia] {url}")
            continue

        results.append((url, page_text))
        print(f"[Wikipedia] pobrano: {url}")
        time.sleep(DELAY)
    return results


def crawl():
    visited = set()
    visited_files = set()  # osobny zbior, zeby nie pobierac tego samego pliku dwa razy
    queue = deque(START_URLS)
    count = 0
    pages = []

    while queue:
        if MAX_PAGES and count >= MAX_PAGES:
            break

        url = queue.popleft()
        url = clean_url(url)
        if url in visited:
            continue
        visited.add(url)

        if should_skip(url):
            continue

        try:
            resp = requests.get(url, headers=HEADERS, timeout=15)
            resp.raise_for_status()
        except Exception as e:
            print(f"[BŁĄD] {url} -> {e}")
            continue

        content_type = resp.headers.get("Content-Type", "")

        if "text/html" not in content_type:
            ext = detect_file_ext(url, content_type)
            if ext and is_whitelisted_file(url):
                text_extractors = {
                    ".pdf": extract_pdf_bytes,
                    ".docx": extract_docx_bytes,
                    ".txt": extract_txt_bytes,
                }
                file_text = text_extractors[ext](resp.content)
                if file_text is not None:
                    pages.append((url, file_text))
                    count += 1
                    SCRAPED_FILES.append(url)
                    print(f"[{count}] plik (bez rozszerzenia w URL): {url}")
                else:
                    FALLBACK_FILES_USED.append(url)
            continue

        soup = BeautifulSoup(resp.text, "html.parser")
        page_text = extract_text(soup)

        if page_text is None:
            FALLBACK_USED.append(url)
            print(f"[POMINIĘTO - brak main-content] {url}")
        else:
            pages.append((url, page_text))
            count += 1
            print(f"[{count}] pobrano: {url}")

        for a in soup.find_all("a", href=True):
            link = urljoin(url, a["href"])
            link = clean_url(link)
            if not is_same_domain(link) or link in visited or should_skip(link):
                continue

            if is_file_link(link):
                # pliki (pdf/docx/txt) nie trafiaja do kolejki do crawlowania,
                # tylko sa od razu probowane do pobrania (jesli sa na whiteliscie)
                if link in visited_files:
                    continue
                visited_files.add(link)
                file_text = scrape_file(link)
                if file_text is not None:
                    pages.append((link, file_text))
                    count += 1
                    print(f"[{count}] plik: {link}")
                time.sleep(DELAY)
            else:
                queue.append(link)

        time.sleep(DELAY)

    wiki_pages = scrape_wikipedia()

    pages = remove_boilerplate(pages)
    pages = dedupe_content_across_pages(pages)
    pages = pages + wiki_pages

    pages, dropped_hubs = drop_thin_hub_pages(pages)

    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
        for url, text in pages:
            if not text.strip():
                continue
            # oddzielanie stron
            f.write(f"\n\nURL: {url}\n\n")
            f.write(text)

    print(f"\nZapisano {len(pages)} stron do pliku: {OUTPUT_FILE}")
    print(f"Zescrapowano {len(SCRAPED_FILES)} plików (pdf/docx/txt):")
    for u in SCRAPED_FILES:
        print(f"  - {u}")
    print(f"Pominięto (brak main-content) {len(FALLBACK_USED)} stron:")
    for u in FALLBACK_USED:
        print(f"  - {u}")
    print(f"Pominięto (brak treści) {len(FALLBACK_FILES_USED)} plików:")
    for u in FALLBACK_FILES_USED:
        print(f"  - {u}")
    print(f"Pominięto jako puste huby/listingi po deduplikacji ({len(dropped_hubs)}):")
    for u, wc in dropped_hubs:
        print(f"  - {u} ({wc} słów)")


if __name__ == "__main__":
    crawl()